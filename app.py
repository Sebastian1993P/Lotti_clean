import os
import re
import uuid
from io import BytesIO
from datetime import datetime
from collections import defaultdict

from flask import (
    Flask, render_template_string, request, redirect,
    url_for, session, send_file, send_from_directory
)
from werkzeug.utils import secure_filename

# (opcjonalne) torch – nie wymagane do działania
try:
    import torch
    DEVICE = torch.device("cuda" if torch.cuda.is_available() else "cpu")
    print(f"Używane urządzenie: {DEVICE}")
except Exception:
    print("Używane urządzenie: cpu (torch niedostępny)")

# NLP (opcjonalne)
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

# DOCX
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# DB
from flask_sqlalchemy import SQLAlchemy
from werkzeug.security import generate_password_hash, check_password_hash

# -----------------------------------------------------------------------------
# FLASK APP
# -----------------------------------------------------------------------------
app = Flask(__name__, static_folder="static")
app.secret_key = os.environ.get("SECRET_KEY", "change-me")
app.config["SQLALCHEMY_DATABASE_URI"] = "sqlite:///lotti.db"
app.config["SQLALCHEMY_TRACK_MODIFICATIONS"] = False
db = SQLAlchemy(app)

# -----------------------------------------------------------------------------
# MODELE
# -----------------------------------------------------------------------------
class User(db.Model):
    __tablename__ = 'user'
    id            = db.Column(db.Integer, primary_key=True)
    username      = db.Column(db.String(64), unique=True, nullable=False)
    password_hash = db.Column(db.String(128), nullable=False)
    cabinets      = db.relationship('Cabinet', backref='owner', lazy=True)

    def set_password(self, pwd):
        self.password_hash = generate_password_hash(pwd)

    def check_password(self, pwd):
        return check_password_hash(self.password_hash, pwd)


class TreatmentType(db.Model):
    __tablename__ = 'treatment_type'
    id                  = db.Column(db.Integer, primary_key=True)
    name                = db.Column(db.String(64), unique=True, nullable=False)
    default_price       = db.Column(db.Float, nullable=False)
    default_description = db.Column(db.Text,   nullable=True)


class ProcedureCode(db.Model):
    __tablename__ = 'procedure_code'
    code             = db.Column(db.String(16), primary_key=True)
    category_name    = db.Column(db.String(64), db.ForeignKey('treatment_type.name'), nullable=False)
    default_duration = db.Column(db.Integer, nullable=False)


class Cabinet(db.Model):
    __tablename__ = 'cabinet'
    id           = db.Column(db.Integer, primary_key=True)
    name         = db.Column(db.String(128), nullable=False)
    logo         = db.Column(db.String(256))
    doctor_name  = db.Column(db.String(128))
    street       = db.Column(db.String(128))
    flat_number  = db.Column(db.String(16))
    postal_code  = db.Column(db.String(16))
    city         = db.Column(db.String(64))
    user_id      = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    treatments   = db.relationship('Treatment', backref='cabinet', lazy=True)


class Treatment(db.Model):
    __tablename__ = 'treatment'
    id              = db.Column(db.Integer, primary_key=True)
    cabinet_id      = db.Column(db.Integer, db.ForeignKey('cabinet.id'), nullable=False)
    type            = db.Column(db.String(64), nullable=False)
    description     = db.Column(db.Text)
    price           = db.Column(db.Float)
    duration        = db.Column(db.Integer, nullable=True)  # ogólny czas (min)
    base_price      = db.Column(db.Float, nullable=True)    # gingiwoplastyka: cena podstawowa
    per_tooth_price = db.Column(db.Float, nullable=True)    # gingiwoplastyka: cena za ząb


class CodeDuration(db.Model):
    __tablename__ = 'code_duration'
    id              = db.Column(db.Integer, primary_key=True)
    cabinet_id      = db.Column(db.Integer, db.ForeignKey('cabinet.id'), nullable=False)
    procedure_code  = db.Column(db.String(16), db.ForeignKey('procedure_code.code'), nullable=False)
    duration        = db.Column(db.Integer, nullable=False)
    timestamp       = db.Column(db.DateTime, default=datetime.utcnow, nullable=False)

    cabinet = db.relationship('Cabinet', backref='code_durations')
    code    = db.relationship('ProcedureCode')


class GeneratedPlan(db.Model):
    __tablename__ = "generated_plan"
    id         = db.Column(db.Integer, primary_key=True)
    user_id    = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    cabinet_id = db.Column(db.Integer, db.ForeignKey('cabinet.id'), nullable=False)
    input_data = db.Column(db.Text,   nullable=False)
    plan_text  = db.Column(db.Text,   nullable=False)
    created_at = db.Column(db.DateTime, default=datetime.utcnow, nullable=False)

    user    = db.relationship('User', backref='generated_plans')
    cabinet = db.relationship('Cabinet', backref='generated_plans')

# -----------------------------------------------------------------------------
# STATIC / UPLOADS
# -----------------------------------------------------------------------------
UPLOAD_FOLDER = os.path.join(app.static_folder, 'uploads')
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# -----------------------------------------------------------------------------
# SZABLONY (pełne)
# -----------------------------------------------------------------------------
login_template = """
<!doctype html>
<html lang="pl">
<head>
  <meta charset="utf-8">
  <title>Panel Logowania</title>
  <link href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.0/dist/css/bootstrap.min.css" rel="stylesheet">
</head>
<body class="p-4">
  <div class="container" style="max-width:480px">
    <div class="text-center mb-4">
      <img src="{{ url_for('static', filename='Lottiimage.png') }}" alt="Logo" style="max-width:200px; margin-bottom: 20px;">
      <h1 class="h4">Logowanie do systemu</h1>
    </div>
    {% if error %}<div class="alert alert-danger">{{ error }}</div>{% endif %}
    <form method="POST" action="/login" class="card card-body">
      <label class="form-label">Nazwa użytkownika</label>
      <input class="form-control mb-3" type="text" name="username" required>
      <label class="form-label">Hasło</label>
      <input class="form-control mb-3" type="password" name="password" required>
      <button class="btn btn-primary w-100">Zaloguj</button>
    </form>
  </div>
</body>
</html>
"""

cabinets_template = """
<!doctype html>
<html lang="pl">
<head>
  <meta charset="utf-8">
  <title>AdminLotti – Gabinety</title>

  <!-- Bootstrap CSS -->
  <link
    href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.0/dist/css/bootstrap.min.css"
    rel="stylesheet">

  <!-- FontAwesome -->
  <link
    href="https://cdnjs.cloudflare.com/ajax/libs/font-awesome/6.1.1/css/all.min.css"
    rel="stylesheet">

  <style>
    body {
      background: #f0f4f8;
      color: #33475b;
      padding-top: 1rem;
      padding-bottom: 2rem;
    }
    .navbar {
      background: #ffffffcc;
      box-shadow: 0 2px 8px rgba(0,0,0,0.05);
      backdrop-filter: blur(10px);
    }
    .navbar-brand, .nav-link {
      color: #33475b !important;
      font-weight: 500;
    }
    .btn-rounded {
      border-radius: 50px;
      padding: .5rem 1.25rem;
      transition: all .2s ease-in-out;
    }
    .btn-rounded:hover {
      transform: translateY(-2px);
    }
    .card-custom {
      border: none;
      border-radius: .75rem;
      box-shadow: 0 2px 10px rgba(0,0,0,0.05);
    }
  </style>
</head>

<body>
  <!-- Navbar -->
  <nav class="navbar navbar-expand-lg sticky-top mb-4">
    <div class="container">
      <a class="navbar-brand" href="/"><i class="fa-solid fa-tooth me-2"></i>ChatLotti</a>
      <div class="ms-auto">
        <a class="nav-link d-inline px-3" href="/"><i class="fa-solid fa-calendar-check me-1"></i>Generuj plan</a>
        <a class="nav-link d-inline px-3" href="/logout"><i class="fa-solid fa-right-from-bracket me-1"></i>Wyloguj</a>
      </div>
    </div>
  </nav>

  <div class="container">
    <!-- Lista gabinetów -->
    <div class="card card-custom mb-4">
      <div class="card-body">
        <h5 class="card-title"><i class="fa-solid fa-building me-2"></i>Zarządzanie gabinetami</h5>
        {% if message %}
          <div class="alert alert-success">{{ message }}</div>
        {% endif %}

        <table class="table align-middle">
          <thead class="table-light">
            <tr>
              <th>Nazwa</th>
              <th>Logo</th>
              <th>Akcje</th>
            </tr>
          </thead>
          <tbody>
            {% for c in cabinets %}
              <tr>
                <td>{{ c.name }}</td>
                <td>
                  {% if c.logo %}
                    <img src="{{ url_for('static', filename='uploads/' ~ c.logo) }}"
                         class="img-fluid rounded" style="max-height:40px;">
                  {% endif %}
                </td>
                <td>
                  <a class="btn btn-sm btn-outline-primary btn-rounded"
                     href="{{ url_for('admin_treatments', cabinet_id=c.id) }}">
                    <i class="fa-solid fa-tooth me-1"></i>Zabiegi
                  </a>
                </td>
              </tr>
            {% endfor %}
          </tbody>
        </table>
      </div>
    </div>

    <!-- Formularz dodawania gabinetu -->
    <div class="card card-custom">
      <div class="card-body">
        <h5 class="card-title"><i class="fa-solid fa-plus me-2"></i>Dodaj nowy gabinet</h5>
        <form method="post" enctype="multipart/form-data">
          <div class="row g-3">
            <div class="col-md-6">
              <label class="form-label">Nazwa gabinetu</label>
              <input type="text" name="name" class="form-control" required>
            </div>
            <div class="col-md-6">
              <label class="form-label">Logo gabinetu</label>
              <input type="file" name="logo" class="form-control" accept="image/*">
            </div>
            <div class="col-md-6">
              <label class="form-label">Imię i nazwisko lekarza</label>
              <input type="text" name="doctor_name" class="form-control" required>
            </div>
            <div class="col-md-6">
              <label class="form-label">Ulica</label>
              <input type="text" name="street" class="form-control" required>
            </div>
            <div class="col-md-3">
              <label class="form-label">Numer lokalu</label>
              <input type="text" name="flat_number" class="form-control">
            </div>
            <div class="col-md-3">
              <label class="form-label">Kod pocztowy</label>
              <input type="text" name="postal_code" pattern="\\d{2}-\\d{3}"
                     class="form-control" placeholder="00-000" required>
            </div>
            <div class="col-md-6">
              <label class="form-label">Miejscowość</label>
              <input type="text" name="city" class="form-control" required>
            </div>
          </div>
          <div class="mt-4">
            <button type="submit" class="btn btn-primary btn-rounded">
              <i class="fa-solid fa-plus me-1"></i>Dodaj gabinet
            </button>
          </div>
        </form>
      </div>
    </div>
  </div>

  <!-- Bootstrap JS Bundle -->
  <script src="https://cdn.jsdelivr.net/npm/bootstrap@5.3.0/dist/js/bootstrap.bundle.min.js"></script>
</body>
</html>
"""

treatments_template = """
<!doctype html>
<html lang="pl">
<head>
  <meta charset="utf-8">
  <title>Zabiegi gabinetu</title>

  <!-- Bootstrap CSS -->
  <link
    href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.0/dist/css/bootstrap.min.css"
    rel="stylesheet">

  <!-- FontAwesome -->
  <link
    href="https://cdnjs.cloudflare.com/ajax/libs/font-awesome/6.1.1/css/all.min.css"
    rel="stylesheet">

  <style>
    body {
      background: #f0f4f8;
      color: #33475b;
      padding-top: 1rem;
      padding-bottom: 2rem;
    }
    .navbar {
      background: #ffffffcc;
      box-shadow: 0 2px 8px rgba(0,0,0,0.05);
      backdrop-filter: blur(10px);
    }
    .navbar-brand, .nav-link {
      color: #33475b !important;
      font-weight: 500;
    }
    .btn-rounded {
      border-radius: 50px;
      padding: .5rem 1.25rem;
      transition: all .2s ease-in-out;
    }
    .btn-rounded:hover {
      transform: translateY(-2px);
    }
    .card-custom {
      border: none;
      border-radius: .75rem;
      box-shadow: 0 2px 10px rgba(0,0,0,0.05);
    }
  </style>
</head>

<body>
<!-- Navbar -->
<nav class="navbar navbar-expand-lg sticky-top mb-4">
  <div class="container">
    <a class="navbar-brand" href="/"><i class="fa-solid fa-tooth me-2"></i>ChatLotti</a>
    <div class="ms-auto">
      <a class="nav-link d-inline px-3" href="/plans"><i class="fa-solid fa-file-alt me-1"></i>Wygenerowane plany</a>
      <a class="nav-link d-inline px-3" href="/"><i class="fa-solid fa-calendar-check me-1"></i>Generuj plan</a>
      <a class="nav-link d-inline px-3" href="/admin/cabinets"><i class="fa-solid fa-hospital me-1"></i>Gabinety</a>
      <a class="nav-link d-inline px-3" href="/logout"><i class="fa-solid fa-right-from-bracket me-1"></i>Wyloguj</a>
    </div>
  </div>
</nav>


  <div class="container">
    <!-- Lista zabiegów -->
    <div class="card card-custom mb-4">
      <div class="card-body">
        <h5 class="card-title"><i class="fa-solid fa-list me-2"></i>Zabiegi w gabinecie &ldquo;{{ cabinet.name }}&rdquo;</h5>
        {% if message %}
          <div class="alert alert-success">{{ message }}</div>
        {% endif %}

        <table class="table align-middle">
          <thead class="table-light">
            <tr>
              <th>Rodzaj</th>
              <th>Opis</th>
              <th>Cena</th>
              <th>Akcje</th>
              <th>Usuń</th>
            </tr>
          </thead>
          <tbody>
            {% for t in treatments %}
              <tr>
                <td>{{ t.type }}</td>
                <td>{{ t.description }}</td>
                <td>
                  {% if t.type == "Gingiwoplastyka" %}
                    {{ t.base_price }} zł + {{ t.per_tooth_price }} zł/ząb
                  {% else %}
                    {{ t.price }} zł
                  {% endif %}
                </td>
                <td>
                  <a class="btn btn-sm btn-outline-primary btn-rounded me-2"
                     href="{{ url_for('edit_treatment', cabinet_id=cabinet.id, treatment_id=t.id) }}">
                    <i class="fa-solid fa-pen-to-square me-1"></i>Edytuj
                  </a>
                  <a class="btn btn-sm btn-outline-secondary btn-rounded"
                     href="{{ url_for('add_duration', cabinet_id=cabinet.id, treatment_id=t.id) }}">
                    <i class="fa-solid fa-clock me-1"></i>Dodaj czas
                  </a>
                </td>
                <td>
                  <form method="POST"
                        action="{{ url_for('delete_treatment', cabinet_id=cabinet.id, treatment_id=t.id) }}"
                        onsubmit="return confirm('Czy na pewno chcesz usunąć ten zabieg?');">
                    <button type="submit" class="btn btn-sm btn-outline-danger btn-rounded">
                      <i class="fa-solid fa-trash me-1"></i>Usuń
                    </button>
                  </form>
                </td>
              </tr>
            {% endfor %}
          </tbody>
        </table>
      </div>
    </div>

    <!-- Formularz dodawania zabiegu -->
    <div class="card card-custom">
      <div class="card-body">
        <h5 class="card-title"><i class="fa-solid fa-plus me-2"></i>Dodaj zabieg</h5>
        <form method="post">
          <div class="mb-3">
            <label class="form-label">Rodzaj</label>
            <select name="type" class="form-select" required>
              {% for tt in types %}
                <option value="{{ tt.name }}">{{ tt.name }}</option>
              {% endfor %}
            </select>
          </div>
          <div class="mb-3">
            <label class="form-label">Opis</label>
            <textarea name="description" class="form-control" rows="2"></textarea>
          </div>
            <!-- wspólne pola -->
          <div class="mb-3" id="price-field">
            <label class="form-label">Cena (PLN)</label>
            <input type="number" name="price" class="form-control" required>
          </div>

          <!-- dodatkowe pola tylko dla Gingiwoplastyka -->
          <div id="gingi-prices" style="display:none;">
            <div class="mb-3">
              <label class="form-label">Cena podstawowa (PLN)</label>
              <input type="number" step="0.01" name="base_price" class="form-control">
            </div>
            <div class="mb-3">
              <label class="form-label">Cena za 1 ząb (PLN)</label>
              <input id="per-tooth-price" type="number" step="0.01" name="per_tooth_price" class="form-control">
            </div>
          </div>

          <script>
            const typeSelect    = document.querySelector('select[name="type"]');
            const priceField    = document.getElementById('price-field');
            const gingiPrices   = document.getElementById('gingi-prices');
            const priceInput    = document.querySelector('input[name="price"]');
            const baseInput     = document.querySelector('input[name="base_price"]');
            const perToothInput = document.getElementById('per-tooth-price');

            function toggleFields() {
              const isGingi = typeSelect.value === 'Gingiwoplastyka';
              priceField.style.display    = isGingi ? 'none' : 'block';
              gingiPrices.style.display   = isGingi ? 'block' : 'none';
              priceInput.required         = !isGingi;
              baseInput.required          = isGingi;
              perToothInput.required      = isGingi;
            }

            typeSelect.addEventListener('change', toggleFields);
            toggleFields();
          </script>

          <button type="submit" class="btn btn-primary btn-rounded">
            <i class="fa-solid fa-plus me-1"></i>Dodaj zabieg
          </button>
        </form>
      </div>
    </div>
  </div>

  <!-- Bootstrap JS Bundle -->
  <script src="https://cdn.jsdelivr.net/npm/bootstrap@5.3.0/dist/js/bootstrap.bundle.min.js"></script>
</body>
</html>
"""

edit_treatment_template = """
<!doctype html>
<html lang="pl">
<head>
  <meta charset="utf-8">
  <title>Edytuj zabieg</title>

  <!-- Bootstrap CSS -->
  <link
    href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.0/dist/css/bootstrap.min.css"
    rel="stylesheet">

  <!-- FontAwesome -->
  <link
    href="https://cdnjs.cloudflare.com/ajax/libs/font-awesome/6.1.1/css/all.min.css"
    rel="stylesheet">

  <style>
    body {
      background: #f0f4f8;
      color: #33475b;
      padding-top: 1rem;
      padding-bottom: 2rem;
    }
    .navbar {
      background: #ffffffcc;
      box-shadow: 0 2px 8px rgba(0,0,0,0.05);
      backdrop-filter: blur(10px);
    }
    .navbar-brand, .nav-link {
      color: #33475b !important;
      font-weight: 500;
    }
    .btn-rounded {
      border-radius: 50px;
      padding: .5rem 1.25rem;
      transition: all .2s ease-in-out;
    }
    .btn-rounded:hover {
      transform: translateY(-2px);
    }
    .card-custom {
      border: none;
      border-radius: .75rem;
      box-shadow: 0 2px 10px rgba(0,0,0,0.05);
    }
    .form-label {
      font-weight: 500;
    }
  </style>
</head>

<body>
 <!-- Navbar -->
<nav class="navbar navbar-expand-lg sticky-top mb-4">
  <div class="container">
    <a class="navbar-brand" href="/"><i class="fa-solid fa-tooth me-2"></i>ChatLotti</a>
    <div class="ms-auto">
      <a class="nav-link d-inline px-3" href="/plans"><i class="fa-solid fa-file-alt me-1"></i>Wygenerowane plany</a>
      <a class="nav-link d-inline px-3" href="/"><i class="fa-solid fa-calendar-check me-1"></i>Generuj plan</a>
      <a class="nav-link d-inline px-3" href="/admin/cabinets"><i class="fa-solid fa-hospital me-1"></i>Gabinety</a>
      <a class="nav-link d-inline px-3" href="/logout"><i class="fa-solid fa-right-from-bracket me-1"></i>Wyloguj</a>
    </div>
  </div>
</nav>


  <div class="container">
    <div class="card card-custom mx-auto" style="max-width: 600px;">
      <div class="card-body">
        <h5 class="card-title mb-4"><i class="fa-solid fa-pen-to-square me-2"></i>Edytuj zabieg „{{ treatment.type }}”</h5>
        <form method="post">
          <div class="mb-3">
            <label class="form-label">Rodzaj (nieedytowalny):</label>
            <input type="text"
                   class="form-control"
                   value="{{ treatment.type }}"
                   disabled>
          </div>
          <div class="mb-3">
            <label class="form-label">Opis:</label>
            <textarea name="description"
                      class="form-control"
                      rows="2">{{ treatment.description }}</textarea>
          </div>
          <div class="mb-3">
            <label class="form-label">Cena (PLN):</label>
            <input type="number"
                   name="price"
                   class="form-control"
                   value="{{ treatment.price }}"
                   required>
          </div>
          <div class="d-flex justify-content-between">
            <a href="{{ url_for('admin_treatments', cabinet_id=cabinet.id) }}"
               class="btn btn-outline-secondary btn-rounded">
              <i class="fa-solid fa-arrow-left me-1"></i>Powrót
            </a>
            <button type="submit"
                    class="btn btn-primary btn-rounded">
              <i class="fa-solid fa-check me-1"></i>Zapisz zmiany
            </button>
          </div>
        </form>
      </div>
    </div>
  </div>

  <!-- Bootstrap JS Bundle -->
  <script src="https://cdn.jsdelivr.net/npm/bootstrap@5.3.0/dist/js/bootstrap.bundle.min.js"></script>
</body>
</html>
"""
durations_template = """
<!doctype html>
<html lang="pl">
<head>
  <meta charset="utf-8">
  <title>Dodaj czas wizyty</title>

  <!-- Bootstrap CSS -->
  <link
    href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.0/dist/css/bootstrap.min.css"
    rel="stylesheet">

  <!-- FontAwesome -->
  <link
    href="https://cdnjs.cloudflare.com/ajax/libs/font-awesome/6.1.1/css/all.min.css"
    rel="stylesheet">

  <style>
    body {
      background: #f0f4f8;
      color: #33475b;
      padding-top: 1rem;
      padding-bottom: 2rem;
    }
    .navbar {
      background: #ffffffcc;
      box-shadow: 0 2px 8px rgba(0,0,0,0.05);
      backdrop-filter: blur(10px);
    }
    .navbar-brand, .nav-link {
      color: #33475b !important;
      font-weight: 500;
    }
    .card-custom {
      border: none;
      border-radius: .75rem;
      box-shadow: 0 2px 10px rgba(0,0,0,0.05);
    }
    .btn-rounded {
      border-radius: 50px;
      padding: .5rem 1.25rem;
      transition: all .2s ease-in-out;
    }
    .btn-rounded:hover {
      transform: translateY(-2px);
    }

    /* Kontener na tabelę historii – ograniczona wysokość, przewijalność */
    .history-container {
      max-height: 250px;
      overflow-y: auto;
      margin-bottom: 1rem;
    }
    .history-container table {
      margin-bottom: 0;
    }
  </style>
</head>

<body>
  <!-- Navbar -->
  <nav class="navbar navbar-expand-lg sticky-top mb-4">
    <div class="container">
      <a class="navbar-brand" href="/"><i class="fa-solid fa-tooth me-2"></i>ChatLotti</a>
      <div class="ms-auto">
        <a class="nav-link d-inline px-3" href="/admin/cabinets">
          <i class="fa-solid fa-hospital me-1"></i>Gabinety
        </a>
        <a class="nav-link d-inline px-3" href="/logout">
          <i class="fa-solid fa-right-from-bracket me-1"></i>Wyloguj
        </a>
      </div>
    </div>
  </nav>

  <div class="container">
    <div class="card card-custom p-4">
      <h5 class="card-title mb-3">
        <i class="fa-solid fa-clock me-2"></i>Dodaj czas wizyty dla „{{ treatment.type }}”
      </h5>

      <!-- CZĘŚĆ 1: Tabela historii -->
      <div class="history-container">
        <table class="table table-striped table-bordered">
          <thead class="table-light">
            <tr>
              <th scope="col">Rodzaj zabiegu</th>
              <th scope="col">Kod procedury</th>
              <th scope="col">Wszystkie dodane czasy [min]</th>
              <th scope="col">Optymalny czas [min]</th>
            </tr>
          </thead>
          <tbody>
            {% if history_groups %}
              {% for group in history_groups %}
                <tr>
                  <td>{{ group.category }}</td>
                  <td>{{ group.procedure_code }}</td>
                  <td>{{ group.durations | join(', ') }} min</td>
                  <td>{{ group.optimal }} min</td>
                </tr>
              {% endfor %}
            {% else %}
              <tr>
                <td colspan="4" class="text-center text-muted">
                  Brak dotychczas zapisanych czasów dla tego rodzaju zabiegu.
                </td>
              </tr>
            {% endif %}
          </tbody>
        </table>
      </div>
      <!-- KONIEC: Tabela historii -->

      <p>Optymalny czas wizyty (dla bieżącego kodu): <strong>{{ optimal }} min</strong></p>

      <form method="post">
        <div class="mb-3">
          <label class="form-label">Rodzaj zabiegu</label>
          <select class="form-select" disabled>
            <option>{{ treatment.type }}</option>
          </select>
        </div>

        <div class="mb-3">
          <label class="form-label" for="code">Kod zabiegu</label>
          <input
            type="text"
            id="code"
            name="procedure_code"
            list="code-list"
            class="form-control"
            required
            value="{{ code or '' }}"
            placeholder="np. 13 MOD">
          <datalist id="code-list">
            {% for c in codes %}
              <option value="{{ c }}">
            {% endfor %}
          </datalist>
        </div>

        <div id="times" class="mb-3">
          <label class="form-label">Czas wizyty 1 (min)</label>
          <input
            type="number"
            name="new_durations[]"
            class="form-control"
            required
            placeholder="np. 75">
        </div>

        <div class="d-flex gap-2">
          <button
            type="button"
            class="btn btn-outline-secondary btn-rounded"
            onclick="addField()">
            <i class="fa-solid fa-plus me-1"></i>Dodaj kolejny czas
          </button>
          <button
            type="submit"
            class="btn btn-primary btn-rounded">
            Zapisz
          </button>
          <a
            href="{{ url_for('admin_treatments', cabinet_id=cabinet.id) }}"
            class="btn btn-outline-danger btn-rounded">
            Wyjdź
          </a>
        </div>
      </form>
    </div>
  </div>

  <!-- Bootstrap JS Bundle -->
  <script
    src="https://cdn.jsdelivr.net/npm/bootstrap@5.3.0/dist/js/bootstrap.bundle.min.js">
  </script>
  <script>
    function addField() {
      const container = document.getElementById('times');
      const idx = container.querySelectorAll('input').length + 1;
      const div = document.createElement('div');
      div.className = 'mb-3';
      div.innerHTML = `
        <label class="form-label">Czas wizyty ${idx} (min)</label>
        <input type="number" class="form-control" name="new_durations[]" required placeholder="np. 80">
      `;
      container.appendChild(div);
    }
  </script>
</body>
</html>
"""

main_template = """<!doctype html>
<html lang="pl">
<head>
  <meta charset="utf-8">
  <title>ChatLotti – Generowanie planu leczenia</title>

  <!-- 1. Bootstrap CSS -->
  <link
    href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.0/dist/css/bootstrap.min.css"
    rel="stylesheet">

  <!-- 2. Ikony -->
  <link
    href="https://cdnjs.cloudflare.com/ajax/libs/font-awesome/6.1.1/css/all.min.css"
    rel="stylesheet">

  <style>
    body {
      background: #f0f4f8;
      color: #33475b;
      padding-top: 1rem;
      padding-bottom: 2rem;
    }
    .navbar {
      background: #ffffffcc;
      box-shadow: 0 2px 8px rgba(0,0,0,0.05);
      backdrop-filter: blur(10px);
    }
    .navbar-brand, .nav-link {
      color: #33475b !important;
      font-weight: 500;
    }
    .hero {
      background: linear-gradient(135deg, #3b8beb 0%, #6fa4ff 100%);
      color: #fff;
      padding: 2.5rem;
      border-radius: .75rem;
      box-shadow: 0 4px 15px rgba(0,0,0,0.1);
      text-align: center;
      margin-bottom: 2rem;
    }
    .btn-rounded {
      border-radius: 50px;
      padding: .75rem 1.5rem;
      box-shadow: 0 2px 6px rgba(0, 0, 0, 0.15);
      transition: all .2s ease-in-out;
    }
    .btn-rounded:hover {
      transform: translateY(-2px);
      box-shadow: 0 4px 12px rgba(0, 0, 0, 0.2);
    }
    .card-custom {
      border: none;
      border-radius: .75rem;
      box-shadow: 0 2px 10px rgba(0,0,0,0.05);
    }
    pre.plan {
      background: #e9ecef;
      border-radius: .5rem;
      padding: 1rem;
      font-family: monospace;
      white-space: pre-wrap;
    }
  </style>
</head>

<body>
  <!-- Navbar -->
  <nav class="navbar navbar-expand-lg sticky-top mb-4">
    <div class="container">
      <a class="navbar-brand" href="/"><i class="fa-solid fa-tooth me-2"></i>ChatLotti</a>
      <div class="ms-auto">
        <a class="nav-link d-inline px-3" href="/admin/cabinets"><i class="fa-solid fa-hospital me-1"></i> Gabinety</a>
        <a class="nav-link d-inline px-3" href="/logout"><i class="fa-solid fa-right-from-bracket me-1"></i> Wyloguj</a>
      </div>
    </div>
  </nav>

  <!-- Hero -->
  <div class="container">
    <div class="hero">
      <h1 class="display-6"><i class="fa-solid fa-calendar-check me-2"></i>Generowanie planu leczenia</h1>
      <p class="lead">Wprowadź kody zabiegów, a system rozłoży je na optymalne wizyty.</p>
    </div>
  </div>

  <!-- Formularz i Wyniki -->
  <div class="container">
    <div class="row gy-4">
      <!-- Formularz -->
      <div class="col-lg-4">
        <div class="card card-custom p-4">
          <h5 class="mb-3"><i class="fa-solid fa-microscope me-2"></i>Nowy plan</h5>
          {% if not cabinets %}
            <div class="alert alert-warning">Brak gabinetów. Dodaj gabinet w zakładce „Gabinety”.</div>
          {% else %}
            <form method="post">
              <div class="mb-3">
                <label class="form-label" for="cabinet_id">Gabinet</label>
                {% if cabinets|length > 1 %}
                  <select class="form-select" id="cabinet_id" name="cabinet_id" required>
                    <option value="">-- wybierz --</option>
                    {% for c in cabinets %}
                      <option value="{{ c.id }}" {% if c.id==selected_id %}selected{% endif %}>
                        {{ c.name }}
                      </option>
                    {% endfor %}
                  </select>
                {% else %}
                  <input type="hidden" name="cabinet_id" value="{{ cabinets[0].id }}">
                  <p class="form-control-plaintext"><strong>{{ cabinets[0].name }}</strong></p>
                {% endif %}
              </div>
              <div class="mb-3">
                <label class="form-label" for="input_data">Kody zabiegów</label>
                <textarea
                  class="form-control"
                  id="input_data"
                  name="input_data"
                  rows="4"
                  placeholder="np. 13 MOD, 14 MO">{{ input_data }}</textarea>
              </div>
              {% if error %}
                <div class="alert alert-danger">{{ error }}</div>
              {% endif %}
              <button type="submit" class="btn btn-primary btn-rounded w-100">
                <i class="fa-solid fa-magic me-1"></i> Generuj plan
              </button>
            </form>
          {% endif %}
        </div>
      </div>

      <!-- Wyniki -->
      <div class="col-lg-8">
        {% if result %}
          <div class="card card-custom mb-4">
            <div class="card-body">
              <h5 class="card-title"><i class="fa-solid fa-file-lines me-2"></i>Wygenerowany plan leczenia</h5>
              <pre class="plan">
{% for category, data in result_items if category != "Higienizacja" and data.teeth %}
{{ loop.index }}. {{ category }}
{% for tooth, time in data['items'] %}
- {{ tooth }}{% if time %} ({{ time }} min){% endif %}
{% endfor %}

{% if category == "Gingiwoplastyka" %}
Łącznie: {{ data.cost_expr }}
{% elif category == "Konsultacja implantologiczna celem odbudowy braku zęba" %}
Koszt: {{ price_map.get(category,0)|int }} zł
{% else %}
{% set count = data.teeth|length %}
Łącznie {{ count }} x {{ price_map.get(category,0)|int }} zł = {{ count * price_map.get(category,0)|int }} zł
{% endif %}

{% endfor %}
</pre>



    <div class="d-flex gap-2 mt-3">
      <form method="POST" action="/download" style="margin-right:1rem;">
        <input type="hidden" name="input_data" value="{{ input_data }}">
        <input type="hidden" name="cabinet_id"  value="{{ selected_id }}">
        <button class="btn btn-outline-secondary btn-rounded">
          <i class="fa-solid fa-download me-1"></i> Pobierz (.docx)
        </button>
      </form>
    </div>



            </div>
          </div>
        {% endif %}

      {% if visits %}
  <div class="card card-custom mb-4">
    <div class="card-body">
      <h5 class="card-title">
        <i class="fa-solid fa-calendar-days me-2"></i>Harmonogram wizyt
      </h5>
      <ol class="ps-3">
        {% for v in visits %}
          {% set h = v.minutes // 60 %}
          {% set m = v.minutes % 60 %}

          {% if v.category == "Higienizacja" %}
            <li class="mb-3">
              <strong>Wizyta {{ h }}h {{ m }}min</strong> – Higienizacja.
            </li>
          {% else %}
            {% set teeth_list = v.teeth | join(' i ') %}
            {% if v.category == "Odbudowa protetyczna - nakład" %}
              {% set extra_lbl = " (odbudowa tymczasowa)" %}
            {% else %}
              {% set extra_lbl = "" %}
            {% endif %}

            {% if v.extra == 0 and v.count > 1 %}
              {% set cost_desc = v.count ~ " x " ~ v.unit_price ~ " zł = " ~ v.base_cost ~ " zł" %}
            {% else %}
              {% set parts = [] %}
              {% for tooth in v.teeth %}
                {% set parts = parts + ["1 x " ~ v.unit_price ~ " zł"] %}
              {% endfor %}
              {% set cost_desc = parts | join(" + ") ~ " = " ~ v.base_cost ~ " zł" %}
            {% endif %}

            <li class="mb-3">
              <strong>Wizyta {{ h }}h {{ m }}min</strong>
              – leczenie zęba {{ teeth_list }}{{ extra_lbl }}.
              Przewidywany koszt wizyty to {{ cost_desc }}.
            </li>
          {% endif %}
        {% endfor %}
      </ol>
    </div>
  </div>
{% endif %}

  <!-- Bootstrap JS Bundle -->
  <script src="https://cdn.jsdelivr.net/npm/bootstrap@5.3.0/dist/js/bootstrap.bundle.min.js"></script>
</body>
</html>
"""

plans_list_template = """
<!doctype html>
<html lang="pl">
<head>
  <meta charset="utf-8">
  <title>Wygenerowane plany</title>
  <link href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.0/dist/css/bootstrap.min.css" rel="stylesheet">
  <link href="https://cdnjs.cloudflare.com/ajax/libs/font-awesome/6.1.1/css/all.min.css" rel="stylesheet">
  <style>
    body { background: #f0f4f8; color: #33475b; padding-top: 1rem; padding-bottom: 2rem; }
    .navbar { background: #ffffffcc; box-shadow: 0 2px 8px rgba(0,0,0,0.05); backdrop-filter: blur(10px); }
    .btn-rounded { border-radius: 50px; padding: .5rem 1.25rem; transition: all .2s ease-in-out; }
    .btn-rounded:hover { transform: translateY(-2px); }
    .card-custom { border: none; border-radius: .75rem; box-shadow: 0 2px 10px rgba(0,0,0,0.05); }
  </style>
</head>
<body>
  <nav class="navbar navbar-expand-lg sticky-top mb-4">
    <div class="container">
      <a class="navbar-brand" href="/"><i class="fa-solid fa-tooth me-2"></i>ChatLotti</a>
      <div class="ms-auto">
        <a class="nav-link d-inline px-3" href="/plans"><i class="fa-solid fa-file-alt me-1"></i> Wygenerowane plany</a>
        <a class="nav-link d-inline px-3" href="/admin/cabinets"><i class="fa-solid fa-hospital me-1"></i> Gabinety</a>
        <a class="nav-link d-inline px-3" href="/logout"><i class="fa-solid fa-right-from-bracket me-1"></i> Wyloguj</a>
      </div>
    </div>
  </nav>

  <div class="container">
    <div class="card card-custom mb-4">
      <div class="card-body">
        <h5 class="card-title"><i class="fa-solid fa-list me-2"></i>Wygenerowane plany</h5>
        {% if not plans %}
          <div class="alert alert-warning">Brak zapisanych planów.</div>
        {% else %}
          <table class="table table-hover">
            <thead class="table-light">
              <tr>
                <th>#</th>
                <th>Data i godzina</th>
                <th>Gabinet</th>
                <th>Input (skrócone)</th>
                <th>Akcje</th>
              </tr>
            </thead>
            <tbody>
              {% for p in plans %}
                <tr>
                  <td>{{ loop.index }}</td>
                  <td>{{ p.created_at.strftime("%d.%m.%Y %H:%M:%S") }}</td>
                  <td>{{ p.cabinet.name }}</td>
                  <td style="max-width:200px; overflow:hidden; text-overflow:ellipsis; white-space:nowrap;">
                    {{ p.input_data }}
                  </td>
                  <td>
                    <a href="{{ url_for('view_or_edit_plan', plan_id=p.id) }}"
                       class="btn btn-sm btn-primary btn-rounded me-1">
                      <i class="fa-solid fa-eye me-1"></i>Podgląd / Edytuj
                    </a>
                    <form method="POST"
                          action="{{ url_for('delete_plan', plan_id=p.id) }}"
                          style="display:inline;"
                          onsubmit="return confirm('Czy na pewno usunąć ten plan?');">
                      <button type="submit" class="btn btn-sm btn-danger btn-rounded">
                        <i class="fa-solid fa-trash me-1"></i>Usuń
                      </button>
                    </form>
                  </td>
                </tr>
              {% endfor %}
            </tbody>
          </table>
        {% endif %}
      </div>
    </div>
  </div>

  <script src="https://cdn.jsdelivr.net/npm/bootstrap@5.3.0/dist/js/bootstrap.bundle.min.js"></script>
</body>
</html>
"""

plan_detail_template = """
<!doctype html>
<html lang="pl">
<head>
  <meta charset="utf-8">
  <title>Podgląd / Edytuj plan leczenia</title>

  <!-- Bootstrap CSS -->
  <link
    href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.0/dist/css/bootstrap.min.css"
    rel="stylesheet">
  <!-- FontAwesome -->
  <link
    href="https://cdnjs.cloudflare.com/ajax/libs/font-awesome/6.1.1/css/all.min.css"
    rel="stylesheet">
  <style>
    body {
      background: #f0f4f8;
      color: #33475b;
      padding-top: 1rem;
      padding-bottom: 2rem;
    }
    .navbar {
      background: #ffffffcc;
      box-shadow: 0 2px 8px rgba(0,0,0,0.05);
      backdrop-filter: blur(10px);
    }
    .navbar-brand, .nav-link {
      color: #33475b !important;
      font-weight: 500;
    }
    .hero {
      background: linear-gradient(135deg, #3b8beb 0%, #6fa4ff 100%);
      color: #fff;
      padding: 2.5rem;
      border-radius: .75rem;
      box-shadow: 0 4px 15px rgba(0,0,0,0.1);
      text-align: center;
      margin-bottom: 2rem.
    }
    .btn-rounded {
      border-radius: 50px;
      padding: .75rem 1.5rem;
      box-shadow: 0 2px 6px rgba(0, 0, 0, 0.15);
      transition: all .2s ease-in-out;
    }
    .btn-rounded:hover {
      transform: translateY(-2px);
      box-shadow: 0 4px 12px rgba(0, 0, 0, 0.2);
    }
    .card-custom {
      border: none;
      border-radius: .75rem;
      box-shadow: 0 2px 10px rgba(0,0,0,0.05);
    }
    pre.plan {
      background: #e9ecef;
      border-radius: .5rem;
      padding: 1rem;
      font-family: monospace;
      white-space: pre-wrap;
    }
  </style>
</head>

<body>
  <nav class="navbar navbar-expand-lg sticky-top mb-4">
    <div class="container">
      <a class="navbar-brand" href="/"><i class="fa-solid fa-tooth me-2"></i>ChatLotti</a>
      <div class="ms-auto">
        <a class="nav-link d-inline px-3" href="/plans"><i class="fa-solid fa-file-alt me-1"></i>Wygenerowane plany</a>
        <a class="nav-link d-inline px-3" href="/"><i class="fa-solid fa-calendar-check me-1"></i>Generuj plan</a>
        <a class="nav-link d-inline px-3" href="/admin/cabinets"><i class="fa-solid fa-hospital me-1"></i>Gabinety</a>
        <a class="nav-link d-inline px-3" href="/logout"><i class="fa-solid fa-right-from-bracket me-1"></i>Wyloguj</a>
      </div>
    </div>
  </nav>

  <div class="container">
    <div class="row gy-4">
      <div class="col-lg-4">
        <div class="card card-custom p-4">
          <h5 class="mb-3"><i class="fa-solid fa-pen-to-square me-2"></i>Edytuj plan</h5>
          <form method="post" action="/plans/{{ plan.id }}">
            <input type="hidden" name="cabinet_id" value="{{ plan.cabinet_id }}">
            <div class="mb-3">
              <label class="form-label" for="input_data">Kody zabiegów</label>
              <textarea class="form-control" id="input_data" name="input_data" rows="4" required>{{ plan.input_data }}</textarea>
            </div>
            <button type="submit" class="btn btn-primary btn-rounded w-100">
              <i class="fa-solid fa-magic me-1"></i> Zaktualizuj plan
            </button>
          </form>
        </div>
      </div>

      <div class="col-lg-8">
        {% if result %}
          <div class="card card-custom mb-4">
            <div class="card-body">
              <h5 class="card-title"><i class="fa-solid fa-file-lines me-2"></i>Wygenerowany plan leczenia</h5>
              <pre class="plan">1. Higienizacja
{% for category, data in result_items if category != "Higienizacja" and data.teeth %}
{{ loop.index + 1 }}. {{ category }}
{% for tooth, time in data['items'] %}
- {{ tooth }}
{% endfor %}

{% if category == "Gingiwoplastyka" %}
Łącznie: {{ data.cost_expr }}
{% elif category == "Konsultacja implantologiczna celem odbudowy braku zęba" %}
Koszt: {{ price_map.get(category,0)|int }} zł
{% else %}
{% set count = data.teeth|length %}
Łącznie {{ count }} x {{ price_map.get(category,0)|int }} zł = {{ count * price_map.get(category,0)|int }} zł
{% endif %}

{% endfor %}</pre>
            </div>
          </div>
        {% endif %}

        {% if visits %}
          <div class="card card-custom mb-4">
            <div class="card-body">
              <h5 class="card-title">
                <i class="fa-solid fa-calendar-days me-2"></i>Harmonogram wizyt
              </h5>
              <ol class="ps-3">
                {% for v in visits %}
                  {% set h = v.minutes // 60 %}
                  {% set m = v.minutes % 60 %}
                  {% if v.category == "Higienizacja" %}
                    <li class="mb-3">
                      <strong>Wizyta {{ h }}h {{ m }}min</strong> – Higienizacja.
                    </li>
                  {% else %}
                    {% set teeth_list = v.teeth | join(' i ') %}
                    {% if v.category == "Odbudowa protetyczna - nakład" %}
                      {% set extra_lbl = " (odbudowa tymczasowa)" %}
                    {% else %}
                      {% set extra_lbl = "" %}
                    {% endif %}

                    {% if v.extra == 0 and v.count > 1 %}
                      {% set cost_desc = v.count ~ " x " ~ v.unit_price ~ " zł = " ~ v.base_cost ~ " zł" %}
                    {% else %}
                      {% set parts = [] %}
                      {% for tooth in v.teeth %}
                        {% set parts = parts + ["1 x " ~ v.unit_price ~ " zł"] %}
                      {% endfor %}
                      {% set cost_desc = parts | join(" + ") ~ " = " ~ v.base_cost ~ " zł" %}
                    {% endif %}

                    <li class="mb-3">
                      <strong>Wizyta {{ h }}h {{ m }}min</strong>
                      – leczenie zęba {{ teeth_list }}{{ extra_lbl }}.
                      Przewidywany koszt wizyty to {{ cost_desc }}.
                    </li>
                  {% endif %}
                {% endfor %}
              </ol>
            </div>
          </div>
        {% endif %}

        <div class="d-flex gap-2 mt-3">
          <form method="POST" action="/download" style="margin-right:1rem;">
            <input type="hidden" name="cabinet_id" value="{{ plan.cabinet_id }}">
            <input type="hidden" name="input_data" value="{{ plan.input_data }}">
            <button class="btn btn-outline-secondary btn-rounded">
              <i class="fa-solid fa-download me-1"></i> Pobierz (.docx)
            </button>
          </form>
        </div>

      </div>
    </div>
  </div>

  <script src="https://cdn.jsdelivr.net/npm/bootstrap@5.3.0/dist/js/bootstrap.bundle.min.js"></script>
</body>
</html>
"""
# -----------------------------------------------------------------------------
# POMOCNICZE / LOGIKA
# -----------------------------------------------------------------------------
def analyze_treatment_similarity(input_text: str):
    cats = [
        "Mikroskopowe leczenie odtwórcze",
        "Weryfikacja zębów po leczeniu kanałowym",
        "Odbudowa protetyczna - nakład",
        "Konsultacja implantologiczna celem odbudowy braku zęba",
        "Gingiwoplastyka"
    ]
    try:
        vect = TfidfVectorizer()
        std = vect.fit_transform(cats)
        inp = vect.transform([input_text or ""])
        sims = cosine_similarity(inp, std)
        return dict(zip(cats, sims[0]))
    except Exception:
        return {c: 0.0 for c in cats}

def parse_gingi_range(txt: str):
    m = re.search(r'(\d{2})\s*-\s*(\d{2})', txt)
    if not m:
        return []
    a, b = int(m.group(1)), int(m.group(2))
    raw = [f"{i:02d}" for i in range(a, b + 1)]
    return [code for code in raw if re.match(r'^[1-4][1-8]$', code)]

def parse_input(text: str):
    tokens = re.split(r",\s*", (text or "").strip())
    parsed = []
    i = 0
    while i < len(tokens):
        entry = tokens[i].strip()
        if not entry:
            i += 1
            continue
        if entry.lower().startswith("gingi"):
            rest = entry[len("gingi"):].strip()
            gingi_tokens = []
            if rest:
                gingi_tokens.append(rest)
            j = i + 1
            while j < len(tokens):
                cand = tokens[j].strip()
                if re.match(r"^\d{2}\s*-\s*\d{2}$", cand) or re.match(r"^\d{2}$", cand):
                    gingi_tokens.append(cand)
                    j += 1
                else:
                    break
            for tok in gingi_tokens:
                tok = tok.strip()
                if "-" in tok:
                    for code in parse_gingi_range(f"gingi {tok}"):
                        parsed.append({
                            "tooth_code": code,
                            "treatment_code": "Gingiwoplastyka",
                            "procedure_code": f"{code} Gingiwoplastyka"
                        })
                else:
                    code = tok.zfill(2)
                    parsed.append({
                        "tooth_code": code,
                        "treatment_code": "Gingiwoplastyka",
                        "procedure_code": f"{code} Gingiwoplastyka"
                    })
            i = j
            continue
        m = re.match(r"(\d{2})\s*(.+)", entry)
        if m:
            tooth = m.group(1)
            treat = m.group(2).strip()
            parsed.append({
                "tooth_code": tooth,
                "treatment_code": treat,
                "procedure_code": f"{tooth} {treat}"
            })
        i += 1
    return parsed

def classify_entry(entry):
    tooth_code = entry['tooth_code']
    t = entry['treatment_code'].lower()
    if "po endo" in t:
        return ("Weryfikacja zębów po leczeniu kanałowym", tooth_code)
    if "ex" in t:
        return ("Do usunięcia", tooth_code)
    if t == "brak":
        return ("Konsultacja implantologiczna celem odbudowy braku zęba", tooth_code)
    if "gingi" in t:
        return ("Gingiwoplastyka", tooth_code)
    if "nakład" in t or "naklad" in t:
        return ("Odbudowa protetyczna - nakład", tooth_code)
    if "korona" in t:
        return ("Odbudowa protetyczna - korona", tooth_code)
    code_letters = t.replace(" ", "")
    if code_letters.isalpha():
        if len(code_letters) in [1, 2]:
            return ("Mikroskopowe leczenie odtwórcze", tooth_code)
        if len(code_letters) == 3:
            return ("Odbudowa protetyczna - nakład", tooth_code)
    return (None, tooth_code)

def generate_visit_plan(parsed_entries, duration_map, price_map, per_tooth_map):
    visits = []
    idx = 1
    visits.append({
        "idx": idx, "label": f"Wizyta {idx}", "category": "Higienizacja",
        "unit_price": 0, "count": 0, "teeth": [], "minutes": 60, "base_cost": 0, "extra": 0
    })
    idx += 1

    cbct_cats = {"Weryfikacja zębów po leczeniu kanałowym", "Konsultacja implantologiczna celem odbudowy braku zęba"}
    cbct_items = [e for e in parsed_entries if classify_entry(e)[0] in cbct_cats]
    if cbct_items:
        teeth_list = [e['tooth_code'] for e in cbct_items]
        total_time = sum(duration_map.get(e['procedure_code'], 60) for e in cbct_items)
        category_cbct = classify_entry(cbct_items[0])[0]
        unit_price_cbct = price_map.get(category_cbct, 0)
        base_cost_cbct = sum(price_map.get(classify_entry(e)[0], 0) for e in cbct_items)
        visits.append({
            "idx": idx, "label": f"Wizyta {idx}", "category": category_cbct,
            "unit_price": unit_price_cbct, "count": len(teeth_list),
            "teeth": teeth_list, "minutes": total_time, "base_cost": base_cost_cbct, "extra": 400
        })
        idx += 1

    def cluster_by_tooth_neighborhood(entries_same_category):
        upper_order = [f"1{i}" for i in range(8,0,-1)] + [f"2{i}" for i in range(1,9)]
        lower_order = [f"3{i}" for i in range(8,0,-1)] + [f"4{i}" for i in range(1,9)]
        order_map = {c:("upper",i) for i,c in enumerate(upper_order)}
        order_map.update({c:("lower",i) for i,c in enumerate(lower_order)})

        indexed = []
        for e in entries_same_category:
            cod = e["tooth_code"]
            info = order_map.get(cod)
            if info:
                jaw, pos = info
                indexed.append((jaw, pos, e))

        clusters = []
        for jaw_label in ("upper","lower"):
            jaw_list = [(pos, entry) for (jaw,pos,entry) in indexed if jaw==jaw_label]
            if not jaw_list: continue
            jaw_list.sort(key=lambda x: x[0])
            n = len(jaw_list)
            visited = [False]*n
            for i in range(n):
                if visited[i]: continue
                stack=[i]; visited[i]=True; comp=[]
                while stack:
                    u=stack.pop(); comp.append(u)
                    for v in range(n):
                        if not visited[v] and abs(jaw_list[u][0]-jaw_list[v][0])<=2:
                            visited[v]=True; stack.append(v)
                clusters.append([jaw_list[k][1] for k in comp])
        return clusters

    cbct_and_none = cbct_cats.union({None})
    by_cat = {}
    for e in parsed_entries:
        cat, _ = classify_entry(e)
        if cat in cbct_and_none:
            continue
        by_cat.setdefault(cat, []).append(e)

    for category, entries in by_cat.items():
        clusters = cluster_by_tooth_neighborhood(entries)
        for group in clusters:
            group_sorted = sorted(group, key=lambda e: duration_map.get(e['procedure_code'], 60), reverse=True)
            curr_group, curr_time = [], 0
            for e in group_sorted:
                d = duration_map.get(e['procedure_code'], 60)
                if curr_time + d <= 120:
                    curr_group.append(e); curr_time += d
                else:
                    teeth_codes = [x['tooth_code'] for x in curr_group]
                    if category == "Gingiwoplastyka":
                        base = price_map.get(category, 0); per = per_tooth_map.get(category, 0); cnt = len(curr_group)
                        total_cost = base + cnt*per
                        visits.append({
                            "idx": idx, "label": f"Wizyta {idx}", "category": category,
                            "unit_price": per, "count": cnt, "teeth": teeth_codes,
                            "minutes": curr_time, "base_cost": total_cost, "extra": 0
                        })
                    else:
                        unit = price_map.get(category, 0); cnt = len(curr_group)
                        visits.append({
                            "idx": idx, "label": f"Wizyta {idx}", "category": category,
                            "unit_price": unit, "count": cnt, "teeth": teeth_codes,
                            "minutes": curr_time, "base_cost": cnt*unit, "extra": 0
                        })
                    idx += 1
                    curr_group, curr_time = [e], d
            if curr_group:
                teeth_codes = [x['tooth_code'] for x in curr_group]
                if category == "Gingiwoplastyka":
                    base = price_map.get(category, 0); per = per_tooth_map.get(category, 0); cnt = len(curr_group)
                    total_cost = base + cnt*per
                    visits.append({
                        "idx": idx, "label": f"Wizyta {idx}", "category": category,
                        "unit_price": per, "count": cnt, "teeth": teeth_codes,
                        "minutes": curr_time, "base_cost": total_cost, "extra": 0
                    })
                else:
                    unit = price_map.get(category, 0); cnt = len(curr_group)
                    visits.append({
                        "idx": idx, "label": f"Wizyta {idx}", "category": category,
                        "unit_price": unit, "count": cnt, "teeth": teeth_codes,
                        "minutes": curr_time, "base_cost": cnt*unit, "extra": 0
                    })
                idx += 1
    return visits

def aggregate_plan(parsed_entries, price_map, desc_map, duration_map, per_tooth_map):
    plan = {
        "Higienizacja": {"teeth":["wszystkie zęby"], "cost":None, "description":desc_map.get("Higienizacja","")},
        "Mikroskopowe leczenie odtwórcze": {"teeth":[], "cost":0, "description":desc_map.get("Mikroskopowe leczenie odtwórcze",""), "times":[]},
        "Weryfikacja zębów po leczeniu kanałowym": {"teeth":[], "cost":0, "description":desc_map.get("Weryfikacja zębów po leczeniu kanałowym",""), "times":[]},
        "Odbudowa protetyczna - nakład": {"teeth":[], "cost":0, "description":desc_map.get("Odbudowa protetyczna - nakład",""), "times":[]},
        "Odbudowa protetyczna - korona": {"teeth":[], "cost":0, "description":desc_map.get("Odbudowa protetyczna - korona",""), "times":[]},
        "Konsultacja implantologiczna celem odbudowy braku zęba": {"teeth":[], "cost":0, "description":desc_map.get("Konsultacja implantologiczna celem odbudowy braku zęba",""), "times":[]},
        "Do usunięcia": {"teeth":[], "cost":0, "description":desc_map.get("Do usunięcia",""), "times":[]},
        "Gingiwoplastyka": {"teeth":[], "cost":0, "description":desc_map.get("Gingiwoplastyka",""), "times":[]}
    }
    for entry in parsed_entries:
        category, tooth = classify_entry(entry)
        if category is None: continue
        if category not in plan:
            plan[category] = {"teeth":[], "cost":0, "description":desc_map.get(category,""), "times":[]}
        plan[category]["teeth"].append(tooth)
        plan[category]["times"].append(duration_map.get(entry.get("procedure_code")) or duration_map.get(tooth) or None)
        if category != "Gingiwoplastyka":
            plan[category]["cost"] += price_map.get(category, 0)

    gingi_teeth = plan["Gingiwoplastyka"]["teeth"]
    if gingi_teeth:
        base = price_map.get("Gingiwoplastyka", 0)
        per  = per_tooth_map.get("Gingiwoplastyka", 0)
        total = base + len(gingi_teeth)*per
        plan["Gingiwoplastyka"]["cost"] = total
        plan["Gingiwoplastyka"]["cost_expr"] = f"{base} zł + {len(gingi_teeth)} × {per} zł = {total} zł"
    return plan

def generate_treatment_plan(input_str, price_map, desc_map, duration_map, per_tooth_map):
    parsed = parse_input(input_str)
    return aggregate_plan(parsed, price_map, desc_map, duration_map, per_tooth_map)

def tooth_description(code: str) -> str:
    quad_map = {"1":"prawa górna", "2":"lewa górna", "3":"lewa dolna", "4":"prawa dolna"}
    num_map  = {"1":"jedynka","2":"dwójka","3":"trójka","4":"czwórka","5":"piątka","6":"szóstka","7":"siódemka","8":"ósemka"}
    if len(code)!=2: return ""
    return f"{quad_map.get(code[0],'')} {num_map.get(code[1],'')}".strip()

def format_plan_as_text(plan, price_map):
    lines = ["Wygenerowany plan leczenia:", ""]
    idx = 1
    for category, data in plan.items():
        if category == "Higienizacja" or not data["teeth"]:
            continue
        lines.append(f"{idx}. {category}:")
        lines.append("")
        if category == "Gingiwoplastyka":
            for tooth in data["teeth"]:
                desc = tooth_description(tooth)
                lines.append(f"- {tooth}" + (f" ({desc})" if desc else ""))
            lines.append("")
            expr = data.get("cost_expr","")
            if expr:
                base_part, _, rest = expr.partition("+")
                summary = f"Łącznie: {base_part.strip()} (cena podstawowa) +{rest.strip()}"
                lines.append(summary)
                lines.append("")
        else:
            for tooth in data["teeth"]:
                desc = tooth_description(tooth)
                lines.append(f"- {tooth}" + (f" ({desc})" if desc else ""))
            lines.append("")
            if category == "Konsultacja implantologiczna celem odbudowy braku zęba":
                lines.append(f"Koszt: {price_map.get(category,0)} zł")
                lines.append("")
            else:
                cnt = len(data["teeth"]); unit = price_map.get(category,0); total = cnt*unit
                lines.append(f"Łącznie: {cnt} × {unit} zł = {total} zł")
                lines.append("")
        desc_text = (data.get("description") or "").strip()
        if desc_text:
            lines.append(desc_text); lines.append("")
        idx += 1
    return "\n".join(lines)

def create_word_doc(plan_text, clinic):
    doc = Document()
    normal = doc.styles['Normal']
    normal.font.name = 'Century Gothic'
    normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Century Gothic')
    for h in ['Heading 1','Heading 2','Heading 3']:
        st = doc.styles[h]; st.font.name = 'Century Gothic'
        st._element.rPr.rFonts.set(qn('w:eastAsia'), 'Century Gothic')

    if clinic.get('logo_path') and os.path.exists(clinic['logo_path']):
        p_logo = doc.add_paragraph(); p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_logo.add_run().add_picture(clinic['logo_path'], width=Inches(1.5))

    date_para = doc.add_paragraph(); date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_para.add_run(f"{clinic['city']} {datetime.now():%d.%m.%Y} r.").bold = True

    info = (
        f"Lek. dent. {clinic['doctor_name']}\n"
        f"{clinic['clinic_name']}\n"
        f"{clinic['street']} {clinic['flat_number']}\n"
        f"{clinic['postal_code']} {clinic['city']}"
    )
    doc.add_paragraph(info)

    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("WSTĘPNY PLAN LECZENIA"); r.bold = True; r.font.size = Pt(14)
    doc.add_paragraph()

    for line in plan_text.splitlines():
        if re.match(r'^\d+\.\s+.+', line):
            doc.add_heading(line.strip(), level=2)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style='List Bullet')
        else:
            doc.add_paragraph(line)
    f = BytesIO(); doc.save(f); f.seek(0); return f

# -----------------------------------------------------------------------------
# AUTH / GUARD
# -----------------------------------------------------------------------------
@app.before_request
def require_login():
    allowed = {"login", "static", "healthz", "__healthz", "debug_image"}
    if request.endpoint not in allowed and "user_id" not in session:
        return redirect(url_for("login"))

@app.route("/login", methods=["GET","POST"])
def login():
    error = ""
    if request.method == "POST":
        u = User.query.filter_by(username=request.form["username"]).first()
        if u and u.check_password(request.form["password"]):
            session["user_id"] = u.id
            return redirect(url_for("index"))
        error = "Niepoprawna nazwa użytkownika lub hasło."
    return render_template_string(login_template, error=error)

@app.route("/logout")
def logout():
    session.clear()
    return redirect(url_for("login"))

# -----------------------------------------------------------------------------
# ENDPOINTY GŁÓWNE
# -----------------------------------------------------------------------------
@app.route("/", methods=["GET","POST"])
def index():
    all_cabinets = Cabinet.query.filter_by(user_id=session["user_id"]).all()
    input_data  = ""
    result      = {}
    visits      = []
    error       = ""
    selected_id = None

    types     = TreatmentType.query.all()
    price_map = {t.name: t.default_price       for t in types}
    desc_map  = {t.name: t.default_description for t in types}

    duration_map = {pc.code: pc.default_duration for pc in ProcedureCode.query.all()}

    if request.method == "POST":
        selected_id = request.form.get("cabinet_id")
        input_data  = (request.form.get("input_data") or "").strip()

        if not all_cabinets:
            error = "Brak gabinetów. Dodaj gabinet w zakładce Gabinety."
        elif len(all_cabinets) > 1 and not selected_id:
            error = "Proszę wybrać gabinet z listy."
        else:
            if len(all_cabinets) == 1:
                selected_id = all_cabinets[0].id

            treatments = Treatment.query.filter_by(cabinet_id=selected_id).all()
            cds        = CodeDuration.query.filter_by(cabinet_id=selected_id).all()

            by_code = defaultdict(list)
            for cd in cds:
                by_code[cd.procedure_code].append(cd.duration)
            for proc_code, lst in by_code.items():
                ds = sorted(lst); n=len(ds); trim=int(n*0.1)
                trimmed = ds[trim: n-trim] or ds
                optimal = sum(trimmed)//len(trimmed)
                duration_map[proc_code] = optimal

            per_tooth_map = {}
            for t in treatments:
                desc_map[t.type] = t.description or ""
                if t.type == "Gingiwoplastyka":
                    price_map[t.type]     = t.base_price or 0
                    per_tooth_map[t.type] = t.per_tooth_price or 0
                else:
                    price_map[t.type]     = t.price or 0

            if input_data:
                result = generate_treatment_plan(input_data, price_map, desc_map, duration_map, per_tooth_map)
                parsed = parse_input(input_data)
                visits = generate_visit_plan(parsed, duration_map, price_map, per_tooth_map)
                plan_text = format_plan_as_text(result, price_map)

                new_plan = GeneratedPlan(
                    user_id=session["user_id"],
                    cabinet_id=selected_id,
                    input_data=input_data,
                    plan_text=plan_text
                )
                db.session.add(new_plan); db.session.commit()

    result_items = []
    for cat, data in result.items():
        teeth = data.get("teeth", [])
        times = data.get("times", [])
        data["items"] = list(zip(teeth, times))
        result_items.append((cat, data))

    return render_template_string(
        main_template,
        cabinets=all_cabinets,
        input_data=input_data,
        result=result,
        result_items=result_items,
        visits=visits,
        error=error,
        selected_id=selected_id,
        price_map=price_map,
        duration_map=duration_map
    )

@app.route("/plans")
def list_generated_plans():
    user_id = session.get("user_id")
    if not user_id: return redirect(url_for("login"))
    plans = GeneratedPlan.query.filter_by(user_id=user_id).order_by(GeneratedPlan.created_at.desc()).all()
    return render_template_string(plans_list_template, plans=plans)

@app.route("/plans/<int:plan_id>", methods=["GET","POST"])
def view_or_edit_plan(plan_id):
    plan = GeneratedPlan.query.get_or_404(plan_id)
    if plan.user_id != session.get("user_id"):
        return redirect(url_for("list_generated_plans"))

    types = TreatmentType.query.all()
    price_map = {t.name: t.default_price for t in types}
    desc_map  = {t.name: t.default_description for t in types}

    treatments_db = Treatment.query.filter_by(cabinet_id=plan.cabinet_id).all()
    per_tooth_map = {}
    for t in treatments_db:
        desc_map[t.type] = t.description or ""
        if t.type == "Gingiwoplastyka":
            price_map[t.type] = t.base_price or 0
            per_tooth_map[t.type] = t.per_tooth_price or 0
        else:
            price_map[t.type] = t.price or 0

    duration_map = {pc.code: pc.default_duration for pc in ProcedureCode.query.all()}

    if request.method == "POST":
        new_input = (request.form.get("input_data") or "").strip()
        if new_input:
            parsed = parse_input(new_input)
            new_result = generate_treatment_plan(new_input, price_map, desc_map, duration_map, per_tooth_map)
            new_visits = generate_visit_plan(parsed, duration_map, price_map, per_tooth_map)
            new_plan_text = format_plan_as_text(new_result, price_map)
            plan.input_data = new_input
            plan.plan_text  = new_plan_text
            plan.created_at = datetime.utcnow()
            db.session.commit()
        return redirect(url_for("list_generated_plans"))

    parsed = parse_input(plan.input_data)
    result = generate_treatment_plan(plan.input_data, price_map, desc_map, duration_map, per_tooth_map)
    visits = generate_visit_plan(parsed, duration_map, price_map, per_tooth_map)

    result_items = []
    for cat, data in result.items():
        teeth = data.get("teeth", [])
        times = data.get("times", [])
        data["items"] = list(zip(teeth, times))
        result_items.append((cat, data))

    return render_template_string(
        plan_detail_template,
        plan=plan,
        selected_id=plan.cabinet_id,
        input_data=plan.input_data,
        result=result,
        result_items=result_items,
        visits=visits,
        price_map=price_map,
        duration_map=duration_map
    )

@app.route("/plans/<int:plan_id>/delete", methods=["POST"])
def delete_plan(plan_id):
    plan = GeneratedPlan.query.get_or_404(plan_id)
    if plan.user_id != session.get("user_id"):
        return redirect(url_for("list_generated_plans"))
    db.session.delete(plan); db.session.commit()
    return redirect(url_for("list_generated_plans"))

@app.route("/download", methods=["POST"])
def download_docx():
    cabinet_id = request.form["cabinet_id"]
    input_data = (request.form.get("input_data") or "").strip()
    cabinet    = Cabinet.query.get_or_404(cabinet_id)

    types         = TreatmentType.query.all()
    price_map     = {t.name: t.default_price       for t in types}
    desc_map      = {t.name: t.default_description for t in types}
    treatments    = Treatment.query.filter_by(cabinet_id=cabinet_id).all()
    per_tooth_map = {}
    for t in treatments:
        desc_map[t.type] = t.description
        if t.type == "Gingiwoplastyka":
            price_map[t.type]     = t.base_price or 0
            per_tooth_map[t.type] = t.per_tooth_price or 0
        else:
            price_map[t.type]     = t.price or 0

    duration_map = {pc.code: pc.default_duration for pc in ProcedureCode.query.all()}
    for t in treatments:
        if t.duration is not None:
            duration_map[t.type] = t.duration

    plan      = generate_treatment_plan(input_data, price_map, desc_map, duration_map, per_tooth_map)
    plan_text = format_plan_as_text(plan, price_map)

    logo_path = os.path.join(app.static_folder, "uploads", cabinet.logo) if cabinet.logo \
                else os.path.join(app.static_folder, "Lottiimage.png")

    clinic = {
        "logo_path":   logo_path,
        "doctor_name": cabinet.doctor_name,
        "clinic_name": cabinet.name,
        "street":      cabinet.street,
        "flat_number": cabinet.flat_number,
        "postal_code": cabinet.postal_code,
        "city":        cabinet.city
    }
    f = create_word_doc(plan_text, clinic)
    return send_file(
        f, as_attachment=True, download_name="plan_leczenia.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

@app.route("/admin/cabinets", methods=["GET","POST"])
def admin_cabinets():
    message = ""
    if request.method == "POST":
        name        = request.form["name"].strip()
        logo_file   = request.files.get("logo")
        doctor_name = request.form["doctor_name"].strip()
        street      = request.form["street"].strip()
        flat_number = request.form["flat_number"].strip()
        postal_code = request.form["postal_code"].strip()
        city        = request.form["city"].strip()

        logo_filename = None
        if logo_file and logo_file.filename:
            logo_filename = f"{uuid.uuid4().hex}_{secure_filename(logo_file.filename)}"
            logo_file.save(os.path.join(UPLOAD_FOLDER, logo_filename))

        cab = Cabinet(
            name=name, logo=logo_filename, doctor_name=doctor_name,
            street=street, flat_number=flat_number, postal_code=postal_code,
            city=city, user_id=session["user_id"]
        )
        db.session.add(cab); db.session.commit()
        message = f"Gabinet „{name}” został dodany."

    cabinets = Cabinet.query.filter_by(user_id=session["user_id"]).all()
    return render_template_string(cabinets_template, cabinets=cabinets, message=message)

@app.route("/admin/cabinets/<cabinet_id>/treatments", methods=["GET","POST"])
def admin_treatments(cabinet_id):
    cabinet = Cabinet.query.get_or_404(cabinet_id)
    message = ""
    types = TreatmentType.query.all()

    if request.method == "POST":
        chosen_name = request.form["type"]
        description = request.form["description"].strip()
        tr = Treatment(cabinet_id=cabinet.id, type=chosen_name, description=description)
        if chosen_name == "Gingiwoplastyka":
            tr.base_price      = float(request.form["base_price"])
            tr.per_tooth_price = float(request.form["per_tooth_price"])
        else:
            tr.price = float(request.form["price"])
        db.session.add(tr); db.session.commit()
        message = f"Zabieg „{chosen_name}” dodany."

    treatments = Treatment.query.filter_by(cabinet_id=cabinet.id).all()
    return render_template_string(treatments_template, cabinet=cabinet, treatments=treatments, types=types, message=message)

@app.route("/admin/cabinets/<cabinet_id>/treatments/<treatment_id>/delete", methods=["POST"])
def delete_treatment(cabinet_id, treatment_id):
    cabinet = Cabinet.query.get_or_404(cabinet_id)
    tr = Treatment.query.filter_by(id=treatment_id, cabinet_id=cabinet.id).first_or_404()
    db.session.delete(tr); db.session.commit()
    return redirect(url_for("admin_treatments", cabinet_id=cabinet.id))

@app.route("/admin/cabinets/<cabinet_id>/treatments/<treatment_id>/edit", methods=["GET","POST"])
def edit_treatment(cabinet_id, treatment_id):
    cabinet = Cabinet.query.get_or_404(cabinet_id)
    tr = Treatment.query.filter_by(id=treatment_id, cabinet_id=cabinet.id).first_or_404()
    if request.method == "POST":
        tr.description = request.form["description"].strip()
        tr.price       = float(request.form["price"])
        db.session.commit()
        return redirect(url_for("admin_treatments", cabinet_id=cabinet.id))
    return render_template_string(edit_treatment_template, cabinet=cabinet, treatment=tr)

@app.route("/admin/cabinets/<cabinet_id>/treatments/<treatment_id>/durations", methods=["GET","POST"])
def add_duration(cabinet_id, treatment_id):
    cabinet   = Cabinet.query.get_or_404(cabinet_id)
    treatment = Treatment.query.filter_by(id=treatment_id, cabinet_id=cabinet.id).first_or_404()

    all_durations = (
        CodeDuration.query.filter_by(cabinet_id=cabinet.id)
        .order_by(CodeDuration.timestamp.desc()).all()
    )

    seeded = [pc.code for pc in ProcedureCode.query.filter_by(category_name=treatment.type).all()]
    if not seeded: seeded = [treatment.type]

    history_entries = []
    for entry in all_durations:
        if treatment.type == "Gingiwoplastyka":
            if entry.procedure_code.strip().endswith("Gingiwoplastyka"):
                history_entries.append(entry)
        else:
            m = re.match(r"^(\d{2})\s+(.+)$", entry.procedure_code)
            if m:
                tooth_code = m.group(1); proc_short = m.group(2).strip()
                cat, _ = classify_entry({"tooth_code": tooth_code, "treatment_code": proc_short})
                if cat == treatment.type: history_entries.append(entry)
            else:
                if entry.procedure_code == treatment.type: history_entries.append(entry)

    grouped = defaultdict(list)
    for e in history_entries: grouped[e.procedure_code].append(e)

    history_groups = []
    for proc_code, entries in grouped.items():
        pc_obj = ProcedureCode.query.get(proc_code)
        category_name = pc_obj.category_name if pc_obj else treatment.type
        durations_list_for_code = [e.duration for e in entries]
        ds = sorted(durations_list_for_code); n=len(ds); low=int(n*0.1); high=n-low
        trimmed = ds[low:high] or ds
        optimal_for_this_code = int(sum(trimmed)/len(trimmed))
        history_groups.append({
            "category": category_name, "procedure_code": proc_code,
            "durations": durations_list_for_code, "optimal": optimal_for_this_code
        })
    history_groups.sort(key=lambda x: (x["category"], x["procedure_code"]))

    used = [e.procedure_code for e in history_entries]
    codes = sorted(set(seeded + used))

    code = request.args.get("procedure_code") or request.form.get("procedure_code")
    durations_list = []
    if code:
        durations_list = [d.duration for d in CodeDuration.query.filter_by(
            cabinet_id=cabinet.id, procedure_code=code).all()
        ]

    if request.method == "POST":
        procedure_codes = [c.strip() for c in (request.form.get("procedure_code") or "").split(",") if c.strip()]
        durations_raw = request.form.getlist("new_durations[]")
        durations     = [int(d) for d in durations_raw if str(d).strip()]
        total_time = sum(durations)

        n_codes = len(procedure_codes) or 1
        per_code = total_time // n_codes
        allocs = [per_code]*n_codes
        for pc, allocated in zip(procedure_codes, allocs):
            db.session.add(CodeDuration(
                cabinet_id=cabinet.id, procedure_code=pc, duration=allocated
            ))
        db.session.commit()

        first_code = procedure_codes[0] if procedure_codes else None
        return redirect(url_for('add_duration', cabinet_id=cabinet.id, treatment_id=treatment.id, procedure_code=first_code))

    optimal = "—"
    if durations_list:
        ds = sorted(durations_list); n=len(ds); low=int(n*0.1); high=n-low
        trimmed = ds[low:high] or ds
        optimal = int(sum(trimmed)/len(trimmed))

    return render_template_string(
        durations_template,
        cabinet=cabinet, treatment=treatment, code=code,
        durations=durations_list, optimal=optimal, codes=codes,
        history_groups=history_groups
    )

@app.route("/debug_image")
def debug_image():
    return send_from_directory(app.static_folder, "Lottiimage.png")

@app.route("/healthz")
def healthz():
    return "ok", 200

@app.route("/__healthz")
def __healthz():
    return "", 200

# -----------------------------------------------------------------------------
# SEED BAZY
# -----------------------------------------------------------------------------
with app.app_context():
    db.create_all()
    if not User.query.filter_by(username="admin").first():
        u = User(username="admin"); u.set_password("password"); db.session.add(u); db.session.commit()

    default_types = [
        ("Mikroskopowe leczenie odtwórcze",       "Standardowy opis …", 500),
        ("Weryfikacja zębów po leczeniu kanałowym","Standardowy opis …", 500),
        ("Odbudowa protetyczna - nakład",        "Standardowy opis …", 1900),
        ("Odbudowa protetyczna - korona",        "Standardowy opis …", 2000),
        ("Konsultacja implantologiczna celem odbudowy braku zęba", "Standardowy opis …", 250),
        ("Gingiwoplastyka", "Standardowy opis gingiwoplastyki…", 0),
        ("Higienizacja", "Standardowy opis higienizacji…", 0),
        ("Do usunięcia", "Standardowy opis…", 0),
    ]
    for name, desc, price in default_types:
        if not TreatmentType.query.filter_by(name=name).first():
            db.session.add(TreatmentType(name=name, default_description=desc, default_price=price))
    db.session.commit()

    hist = [
        ("35 OD", 75), ("36 O", 75), ("27 OM", 90),
        ("46 O", 45), ("36 OD", 60), ("46 D", 60), ("17 OM", 45),
    ]
    for code, mins in hist:
        if not ProcedureCode.query.get(code):
            db.session.add(ProcedureCode(code=code, category_name="Mikroskopowe leczenie odtwórcze", default_duration=mins))
    db.session.commit()

# -----------------------------------------------------------------------------
# WSGI alias dla gunicorna
# -----------------------------------------------------------------------------
application = app
